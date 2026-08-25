#!/usr/bin/env python3
"""Añade numeración de página real (campo PAGE / NUMPAGES) al pie de cada .docx generado por
pandoc, antes de convertirlo a PDF con LibreOffice. Necesario porque pandoc por sí solo no añade
numeración de página, y la especificación exige que los PDF finales tengan numeración real."""
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    run.font.size = None

    def field(instr):
        r = OxmlElement("w:r")
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instr
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r.append(fld_begin)
        paragraph._p.append(r)
        r2 = OxmlElement("w:r")
        r2.append(instr_text)
        paragraph._p.append(r2)
        r3 = OxmlElement("w:r")
        r3.append(fld_sep)
        paragraph._p.append(r3)
        r4 = OxmlElement("w:r")
        r4.append(fld_end)
        paragraph._p.append(r4)

    field("PAGE")
    lit = paragraph.add_run(" / ")
    field("NUMPAGES")


def main(path):
    doc = Document(path)
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = 1  # center
    p.add_run("Página ")
    add_page_number_field(p)
    doc.save(path)
    print(f"Numeración de página añadida a {path}")


if __name__ == "__main__":
    for path in sys.argv[1:]:
        main(path)
